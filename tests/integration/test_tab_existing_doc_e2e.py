"""Live end-to-end: flat heading-1 .docx → ``tab_existing_doc`` → native tabs.

Different code path than ``test_retrofit.py``, which exercises
``retrofit_existing_docx`` with marker-based detection (for docs that
have no headings). This test exercises ``convert_docx_to_tabbed_doc``
(the engine behind the ``gdocs_tab_existing_doc`` MCP tool) with the
heading-1 detection path — the more common case where the source doc
is properly styled.

Acceptance per Part 3 of the v1.1.1 test plan:
    flat fixture .docx with Heading 1 sections → tab_existing_doc
    → assert native tabs with correct titles/icons and warnings == [].
"""
from __future__ import annotations

import io

import pytest

pytestmark = pytest.mark.live


def _build_heading1_docx(section_titles: list[str]) -> bytes:
    """A clean .docx with one Heading 1 per ``section_titles`` + body."""
    from docx import Document
    doc = Document()
    for title in section_titles:
        doc.add_heading(title, level=1)
        doc.add_paragraph(f"body content for {title}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_tab_existing_doc_with_heading1_sections(
    live_creds, tmp_path, created_docs,
):
    """Heading-1-styled docx converts to one native tab per heading,
    no warnings. Guards the convert_docx_to_tabbed_doc happy path."""
    from google_docs_mcp.docx_import import convert_docx_to_tabbed_doc

    section_titles = ["Intro", "Background", "Method", "Results", "Discussion"]
    docx_path = tmp_path / "flat_headings.docx"
    docx_path.write_bytes(_build_heading1_docx(section_titles))

    result = convert_docx_to_tabbed_doc(
        live_creds,
        docx_path=docx_path,
        split_by="heading_1",
        title="test_tab_existing_doc_e2e_heading1",
    )

    if result.get("doc_id"):
        created_docs.append(result["doc_id"])

    # ----- Acceptance assertions -----
    assert result.get("doc_id"), f"no doc_id in result: {result!r}"
    assert result.get("url"), f"no url in result: {result!r}"
    assert result.get("split_strategy_used") == "heading_1"

    tabs = result.get("tabs") or []
    tab_titles = [t.get("title") for t in tabs]

    # One tab per heading. Placeholder may or may not be deleted
    # depending on placeholder_behavior; default is 'delete' which
    # removes the empty Tab-1 placeholder, leaving exactly N tabs.
    for expected in section_titles:
        assert expected in tab_titles, (
            f"expected tab titled {expected!r} not in {tab_titles!r}"
        )

    # No real warnings (info messages about empty placeholder section
    # are allowed — they're cosmetic constraints of the Apps Script
    # path, not actionable bugs).
    warnings = result.get("warnings") or []
    assert warnings == [], (
        f"expected zero actionable warnings, got: {warnings!r}. "
        "Apps Script's structural 'Can't remove the last paragraph' "
        "messages go into result.info, not warnings — if those landed "
        "in warnings, the split logic is wrong."
    )
