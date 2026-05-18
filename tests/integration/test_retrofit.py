"""Retrofit: docx with banner markers → all markers matched.

Guards the run-fragmentation fix from v0.15.1. If text extraction
regresses (stops walking w:sym, <w:tab>, etc.) or normalization
drops smart-quote folding, this test catches it.
"""
from __future__ import annotations

import io

import pytest

pytestmark = pytest.mark.live


def _build_pathological_docx() -> bytes:
    """A .docx with NO Heading 1s and the SAME pathologies real Word writes.

    Covers (in one document, end-to-end through Drive's convert):
      - Fragmented runs (multiple <w:r> mid-phrase — Word does this
        for spell-check tags, language tags, rPr changes)
      - NBSP as the literal U+00A0 character inside <w:t> (how Word
        actually writes it; the ``<w:sym w:char="00A0"/>`` form is
        rarely used in practice AND triggers a Drive convert 500, so
        we don't use it here — that case is covered by the in-process
        unit test instead)
      - Smart quotes (Word's autocorrect output)
      - Em-dash (also autocorrect)
    """
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def add_fragmented(doc, parts):
        p = OxmlElement("w:p")
        for part in parts:
            r = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.text = part
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            r.append(t)
            p.append(r)
        doc.element.body.insert(-1, p)

    doc = Document()
    add_fragmented(doc, ["Sec", "tion", " ", "One"])           # fragmented runs
    doc.add_paragraph("body of section one")
    add_fragmented(doc, ["Section", " ", "Two"])          # NBSP as U+00A0 literal
    doc.add_paragraph("body of section two")
    add_fragmented(doc, ["Smart “Quotes” Three"])    # smart quotes
    doc.add_paragraph("body of section three")
    add_fragmented(doc, ["Em—dash Section Four"])         # em-dash
    doc.add_paragraph("body of section four")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_retrofit_pathological_e2e(live_creds, tmp_path):
    """All 4 pathological markers match end-to-end through upload+convert+restructure."""
    from google_docs_mcp.drive_api import trash_drive_file
    from google_docs_mcp.retrofit import retrofit_existing_docx

    docx_path = tmp_path / "retrofit_live.docx"
    docx_path.write_bytes(_build_pathological_docx())

    # Caller types plain ASCII for the smart-quote / em-dash markers —
    # the normalization layer must reconcile.
    markers = [
        {"marker_text": "Section One",            "tab_title": "One"},
        {"marker_text": "Section Two",            "tab_title": "Two"},
        {"marker_text": 'Smart "Quotes" Three',   "tab_title": "Three"},
        {"marker_text": "Em-dash Section Four",   "tab_title": "Four"},
    ]

    result = retrofit_existing_docx(
        live_creds, markers=markers, docx_path=docx_path,
        title="retrofit_e2e_pathological",
    )
    try:
        info = result.get("retrofit") or {}
        matched = info.get("markers_matched")
        missed = info.get("markers_missed") or []
        assert matched == 4, (
            f"expected all 4 markers matched, got {matched}; missed={missed!r}. "
            "Likely run-fragmentation, NBSP-literal, smart-quote, or em-dash "
            "normalization regression."
        )
        assert missed == [], f"unexpected misses: {missed}"
        assert result.get("doc_id")
        assert len(result.get("tabs") or []) >= 4
    finally:
        if result.get("doc_id"):
            try:
                trash_drive_file(live_creds, result["doc_id"])
            except Exception:
                pass
