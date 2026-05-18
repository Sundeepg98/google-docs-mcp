"""Unit-level tests for retrofit's text extraction + normalization.

Guards the v0.15.1 fix for fragmented OOXML runs, NBSP-via-<w:sym>,
typographic substitution (smart quotes, em-dash), and case mismatch.

Runs entirely in-process via python-docx; no Drive involved. The live
integration test (test_retrofit.py) uses simpler content because
Drive's converter 500s on docs that mix all these pathological
constructs at once — that's a Drive limit, not a retrofit bug. This
unit suite is where the pathological coverage lives.
"""
from __future__ import annotations

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _add_fragmented_para(doc, parts):
    """Build a <w:p> with one <w:r> per element. 'NBSP' → <w:sym>."""
    p = OxmlElement("w:p")
    for part in parts:
        r = OxmlElement("w:r")
        if part == "NBSP":
            sym = OxmlElement("w:sym")
            sym.set(qn("w:char"), "00A0")
            r.append(sym)
        else:
            t = OxmlElement("w:t")
            t.text = part
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            r.append(t)
        p.append(r)
    doc.element.body.insert(-1, p)


def test_inject_matches_fragmented_runs():
    """marker_text crossing <w:r> boundaries still matches."""
    from google_docs_mcp.retrofit import _inject_headings

    doc = Document()
    _add_fragmented_para(doc, ["Sec", "tion", " ", "Banner"])
    doc.add_paragraph("body")

    matched, missed = _inject_headings(
        doc, [{"marker_text": "Section Banner", "tab_title": "X"}]
    )
    assert matched == 1, f"missed: {missed}"
    assert missed == []


def test_inject_matches_nbsp_via_sym():
    """NBSP appearing as <w:sym w:char='00A0'/> must read as a space."""
    from google_docs_mcp.retrofit import _inject_headings

    doc = Document()
    _add_fragmented_para(doc, ["Section", "NBSP", "Two"])
    doc.add_paragraph("body")

    matched, missed = _inject_headings(
        doc, [{"marker_text": "Section Two", "tab_title": "Two"}]
    )
    assert matched == 1, f"missed: {missed}"


def test_inject_folds_smart_quotes_to_ascii():
    """Doc has smart quotes; marker has ASCII; must match."""
    from google_docs_mcp.retrofit import _inject_headings

    doc = Document()
    _add_fragmented_para(doc, ["Smart “Quotes” Banner"])
    doc.add_paragraph("body")

    matched, missed = _inject_headings(
        doc, [{"marker_text": 'Smart "Quotes" Banner', "tab_title": "X"}]
    )
    assert matched == 1, f"missed: {missed}"


def test_inject_folds_em_dash_to_hyphen():
    """Doc has em-dash; marker has hyphen; must match."""
    from google_docs_mcp.retrofit import _inject_headings

    doc = Document()
    _add_fragmented_para(doc, ["Em—dash Banner"])
    doc.add_paragraph("body")

    matched, missed = _inject_headings(
        doc, [{"marker_text": "Em-dash Banner", "tab_title": "X"}]
    )
    assert matched == 1, f"missed: {missed}"


def test_inject_case_insensitive_by_default():
    """Default matching is case-insensitive (Word autocorrects case)."""
    from google_docs_mcp.retrofit import _inject_headings

    doc = Document()
    doc.add_paragraph("SECTION FOUR")

    matched, missed = _inject_headings(
        doc, [{"marker_text": "section four", "tab_title": "X"}]
    )
    assert matched == 1, f"missed: {missed}"


def test_inject_case_sensitive_when_requested():
    """case_sensitive=True respects case differences."""
    from google_docs_mcp.retrofit import _inject_headings

    doc = Document()
    doc.add_paragraph("SECTION FOUR")

    matched, missed = _inject_headings(
        doc,
        [{"marker_text": "section four", "tab_title": "X"}],
        case_sensitive=True,
    )
    assert matched == 0, "should NOT match in case-sensitive mode"
    assert len(missed) == 1
    assert missed[0]["marker_text"] == "section four"


def test_inject_missed_marker_returns_candidate_blocks():
    """A miss must include candidate_blocks so the caller can debug."""
    from google_docs_mcp.retrofit import _inject_headings

    doc = Document()
    doc.add_paragraph("The first paragraph")
    doc.add_paragraph("The second paragraph")

    matched, missed = _inject_headings(
        doc, [{"marker_text": "NONEXISTENT BANNER", "tab_title": "X"}]
    )
    assert matched == 0
    assert len(missed) == 1
    candidates = missed[0]["candidate_blocks"]
    assert isinstance(candidates, list) and len(candidates) >= 2
    # Sanity: the actual paragraph text should be in the candidate list
    assert any("first paragraph" in c for c in candidates)


def test_inject_one_block_per_marker_no_double_claim():
    """Two markers with overlapping text don't both target the same block."""
    from google_docs_mcp.retrofit import _inject_headings

    doc = Document()
    doc.add_paragraph("Section Alpha")
    doc.add_paragraph("Section Beta")

    matched, missed = _inject_headings(
        doc,
        [
            {"marker_text": "Section", "tab_title": "X"},
            {"marker_text": "Section", "tab_title": "Y"},
        ],
    )
    # Both match; each claims a separate block in document order
    assert matched == 2
