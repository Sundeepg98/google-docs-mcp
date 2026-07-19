"""Endpoint-level tests for /api/convert?dry_run=1 (wave-3 S1, R-A).

Drives the PRODUCTION app (build_app + real middleware stack + the REAL
NonceStore) so the load-bearing property - a dry_run does NOT consume the
signed URL, and the SAME URL then performs the real convert - is proven
end to end, not against a mock of the thing under test.

The 5 discriminating proofs from the wave-3 contract:
  1. the plan lists the expected tabs + warnings for a fixture docx,
     including a nested case;
  2. NO doc is created and the converter never runs (write-spy);
  3. the SAME signed URL then performs the real convert (nonce preserved)
     and the real result's echo fields match the plan;
  4. no job_store row is created by the dry_run;
  5. a real convert after a dry_run attaches to nothing dry-run-related.

Plus the fingerprint rebase pins: dry_run must NOT enter fp_params (the
_fingerprint spy is never called) and must NOT create job rows.

Unlike the job-model tests, these POST REAL .docx bytes (built in-memory
with python-docx) because the dry_run parses them locally.
"""
from __future__ import annotations

import io
from unittest.mock import patch
from urllib.parse import parse_qs, urlparse

import pytest
from docx import Document
from starlette.testclient import TestClient

from appscriptly import job_store
from appscriptly.http_server import jobs

_TEST_KEY_BYTES = b"test-signing-key-32-characters-long"
_DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


@pytest.fixture(autouse=True)
def isolated_env(tmp_path, monkeypatch):
    monkeypatch.setenv("GOOGLE_DOCS_USER_STORE_PATH", str(tmp_path / "user_state.db"))
    monkeypatch.setenv("GOOGLE_DOCS_DATA_DIR", str(tmp_path))
    monkeypatch.setenv("TRUSTED_HOSTS", "testserver,localhost")
    monkeypatch.delenv("PUBLIC_BASE_URL", raising=False)
    monkeypatch.delenv("GOOGLE_OAUTH_BASE_URL", raising=False)

    job_store._initialized_paths.clear()
    jobs._TASKS.clear()

    from appscriptly.key_provider import InMemoryKeyProvider, with_key_provider
    with with_key_provider(InMemoryKeyProvider({
        "api_bearer": _TEST_KEY_BYTES,
        "oauth_state": _TEST_KEY_BYTES,
        "oauth_state_enc": _TEST_KEY_BYTES,
        "signed_url": _TEST_KEY_BYTES,
    })):
        yield

    job_store._initialized_paths.clear()
    jobs._TASKS.clear()


@pytest.fixture(autouse=True)
def reset_nonce_store():
    from appscriptly import http_server
    from appscriptly.crypto import NonceStore
    from appscriptly.http_server import _state
    fresh = NonceStore()
    _state._NONCE_STORE = fresh
    http_server._NONCE_STORE = fresh
    yield


def _build_app_under_test():
    from fastmcp import FastMCP
    from appscriptly.http_server import build_app
    return build_app(FastMCP("stub-for-dryrun-tests"))


def _mint_qs(user_id: str = "user-A") -> str:
    from appscriptly.crypto import sign_upload_url
    minted = sign_upload_url(
        base_url="http://testserver/api/convert",
        signing_key=_TEST_KEY_BYTES,
        user_id=user_id,
    )
    return urlparse(minted["url"]).query


def _nonce_of(qs: str) -> str:
    return parse_qs(qs)["nonce"][0]


def _bearer_headers() -> dict:
    return {"authorization": f"Bearer {_TEST_KEY_BYTES.decode('utf-8')}"}


def _creds_patches():
    return (
        patch(
            "appscriptly.http_server.routes.convert.get_credentials_for_user",
            return_value="per-user-creds-sentinel",
        ),
        patch(
            "appscriptly.http_server.routes.convert._resolve_client_config",
            return_value={"web": {"client_id": "X", "client_secret": "Y"}},
        ),
        patch(
            "appscriptly.http_server.routes.convert.load_credentials",
            return_value="operator-creds-sentinel",
        ),
    )


def _docx_bytes(
    h1s: list[str] | None = None,
    *,
    leading: str | None = None,
    nested: dict[str, list[str]] | None = None,
) -> bytes:
    """Build a real .docx.

    ``h1s``: flat Heading-1 titles. ``leading``: a body paragraph BEFORE
    the first heading (the placeholder-veto trigger). ``nested``: ordered
    map of Heading-1 title -> list of its Heading-2 child titles.
    """
    doc = Document()
    if leading:
        doc.add_paragraph(leading)
    if nested is not None:
        for parent, children in nested.items():
            doc.add_heading(parent, level=1)
            doc.add_paragraph(f"{parent} body")
            for child in children:
                doc.add_heading(child, level=2)
                doc.add_paragraph(f"{child} body")
    else:
        for title in h1s or []:
            doc.add_heading(title, level=1)
            doc.add_paragraph(f"{title} body")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _form(content: bytes, name: str = "test.docx"):
    return {"file": (name, content, _DOCX_MIME)}


def _count_job_rows() -> int:
    with job_store._connect() as conn:
        return conn.execute("SELECT COUNT(*) FROM convert_jobs").fetchone()[0]


# ---------------------------------------------------------------------
# Proof 1: plan lists tabs + warnings, including a nested case
# ---------------------------------------------------------------------


def test_dry_run_lists_flat_tabs_and_warnings():
    app = _build_app_under_test()
    # A 60-char H1 -> a truncation warning; a normal H1 -> a clean tab.
    long_title = "L" * 60
    content = _docx_bytes([long_title, "Beta"])

    p1, p2, p3 = _creds_patches()
    with p1, p2, p3, TestClient(app) as client:
        r = client.post(
            f"/api/convert?{_mint_qs()}", files=_form(content), data={"dry_run": "1"},
        )
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["dry_run"] is True
    assert body["split_strategy_used"] == "heading_1"
    assert body["heading1_found"] == 2
    assert [t["title"] for t in body["tabs"]] == [long_title[:50], "Beta"]
    assert all(t["depth"] == 0 for t in body["tabs"])
    assert any("truncated to 50" in w for w in body["warnings"])


def test_dry_run_lists_nested_tabs():
    app = _build_app_under_test()
    content = _docx_bytes(nested={"Parent1": ["Child1a", "Child1b"], "Parent2": []})

    p1, p2, p3 = _creds_patches()
    with p1, p2, p3, TestClient(app) as client:
        r = client.post(
            f"/api/convert?{_mint_qs()}",
            files=_form(content),
            data={"dry_run": "1", "nest_by": "heading_2"},
        )
    assert r.status_code == 200, r.text
    body = r.json()
    # heading1_found counts PARENTS only; tabs_created includes children.
    assert body["heading1_found"] == 2
    assert body["tabs_created"] == 4
    assert body["tabs"] == [
        {"title": "Parent1", "depth": 0},
        {"title": "Child1a", "depth": 1},
        {"title": "Child1b", "depth": 1},
        {"title": "Parent2", "depth": 0},
    ]


def test_dry_run_placeholder_veto_on_leading_content():
    app = _build_app_under_test()
    content = _docx_bytes(["Alpha", "Beta"], leading="Intro before any heading.")

    p1, p2, p3 = _creds_patches()
    with p1, p2, p3, TestClient(app) as client:
        r = client.post(
            f"/api/convert?{_mint_qs()}", files=_form(content), data={"dry_run": "1"},
        )
    body = r.json()
    assert body["placeholder"] == "kept"
    assert body["placeholder_veto"] == "unmoved_content"


def test_dry_run_no_splits_reports_problem():
    app = _build_app_under_test()
    doc = Document()
    doc.add_paragraph("Flat document, no headings at all.")
    buf = io.BytesIO()
    doc.save(buf)

    p1, p2, p3 = _creds_patches()
    with p1, p2, p3, TestClient(app) as client:
        r = client.post(
            f"/api/convert?{_mint_qs()}",
            files=_form(buf.getvalue()),
            data={"dry_run": "1"},
        )
    body = r.json()
    assert body["heading1_found"] == 0
    assert body["tabs"] == []
    assert body["problems"] and "No split points" in body["problems"][0]
    # S1-M2: mirrors the real no-splits early return, which reports
    # placeholder "none" (not "kept").
    assert body["placeholder"] == "none"


# ---------------------------------------------------------------------
# Proofs 2-5: no write, nonce preserved, no row, real convert attaches
# to nothing dry-run-related. The load-bearing end-to-end flow.
# ---------------------------------------------------------------------


def test_dry_run_then_real_convert_on_same_url():
    app = _build_app_under_test()
    content = _docx_bytes(["Alpha", "Beta"])
    qs = _mint_qs()
    nonce = _nonce_of(qs)

    convert_calls = []

    def fake_convert(creds, **kwargs):
        # A realistic converter result whose echo fields match the plan.
        convert_calls.append(kwargs)
        return {
            "doc_id": "REAL1", "url": "https://x",
            "heading1_found": 2, "tabs_created": 2,
            "tabs": [{"title": "Alpha"}, {"title": "Beta"}],
            "split_strategy_used": "heading_1", "warnings": [],
        }

    from appscriptly.http_server import _state

    p1, p2, p3 = _creds_patches()
    with p1, p2, p3, patch(
        "appscriptly.http_server.routes.convert._convert_docx",
        side_effect=fake_convert,
    ) as convert_mock, TestClient(app) as client:
        # --- dry_run ---
        dry = client.post(
            f"/api/convert?{qs}", files=_form(content), data={"dry_run": "1"},
        )
        assert dry.status_code == 200, dry.text
        plan = dry.json()

        # Proof 2: the converter NEVER ran, and no doc id came back.
        convert_mock.assert_not_called()
        assert "doc_id" not in plan
        # Proof 4: no job row was created.
        assert _count_job_rows() == 0
        # Nonce is intact (not consumed) after the dry_run.
        assert nonce not in _state._NONCE_STORE._consumed

        # --- real convert on the SAME (still-valid) URL ---
        real = client.post(f"/api/convert?{qs}", files=_form(content))
        assert real.status_code == 200, real.text
        result = real.json()

    # Proof 3: same URL worked for the real convert (nonce was preserved),
    # the converter ran exactly once, and the echo fields match the plan.
    assert len(convert_calls) == 1
    assert result["doc_id"] == "REAL1"
    assert result["heading1_found"] == plan["heading1_found"]
    assert [t["title"] for t in result["tabs"]] == [t["title"] for t in plan["tabs"]]
    assert result["split_strategy_used"] == plan["split_strategy_used"]
    # Proof 5: the real convert is a FRESH job - it attached to nothing the
    # dry_run left behind.
    assert "attached_to_existing_job" not in result
    # The real convert DID consume the nonce (contrast with the dry_run).
    assert nonce in _state._NONCE_STORE._consumed
    assert _count_job_rows() == 1


def test_dry_run_does_not_reach_fingerprint_or_create_rows():
    """The rebase pin, both halves: dry_run never enters fp_params (the
    _fingerprint helper is never called) and never mints a row."""
    app = _build_app_under_test()
    content = _docx_bytes(["Alpha"])

    p1, p2, p3 = _creds_patches()
    with p1, p2, p3, patch(
        "appscriptly.http_server.routes.convert._fingerprint",
    ) as fp_spy, patch(
        "appscriptly.http_server.routes.convert._convert_docx",
        side_effect=AssertionError("converter must not run for a dry_run"),
    ), TestClient(app) as client:
        r = client.post(
            f"/api/convert?{_mint_qs()}", files=_form(content), data={"dry_run": "1"},
        )
    assert r.status_code == 200, r.text
    fp_spy.assert_not_called()
    assert _count_job_rows() == 0


# ---------------------------------------------------------------------
# dry_run + async, batch, drive_file_id, malformed, invalid value
# ---------------------------------------------------------------------


def test_dry_run_wins_over_async_with_info_note():
    app = _build_app_under_test()
    content = _docx_bytes(["Alpha"])

    p1, p2, p3 = _creds_patches()
    with p1, p2, p3, patch(
        "appscriptly.http_server.routes.convert._convert_docx",
        side_effect=AssertionError("converter must not run for a dry_run"),
    ), TestClient(app) as client:
        r = client.post(
            f"/api/convert?{_mint_qs()}",
            files=_form(content),
            data={"dry_run": "1", "async": "1"},
        )
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["dry_run"] is True
    assert "job_id" not in body
    assert any("async was ignored" in note for note in body["info"])


def _banner_docx_bytes() -> bytes:
    """A styled doc with NO headings - the canonical markers use case."""
    doc = Document()
    doc.add_paragraph("MODULE ONE overview")
    doc.add_paragraph("module one body")
    doc.add_paragraph("MODULE TWO overview")
    doc.add_paragraph("module two body")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_dry_run_with_markers_plans_injected_tabs():
    """S1-B1: dry_run + markers must reflect retrofit's injection (the
    real path routes to _retrofit_docx), not report zero headings with
    advice to use markers. Neither converter entry runs; nothing is
    written; the retrofit echo mirrors the real response shape."""
    import json as _json

    app = _build_app_under_test()
    markers = [
        {"marker_text": "MODULE ONE", "tab_title": "One"},
        {"marker_text": "MODULE TWO", "tab_title": "Two"},
    ]

    p1, p2, p3 = _creds_patches()
    with p1, p2, p3, patch(
        "appscriptly.http_server.routes.convert._convert_docx",
        side_effect=AssertionError("converter must not run for a dry_run"),
    ), patch(
        "appscriptly.http_server.routes.convert._retrofit_docx",
        side_effect=AssertionError("retrofit must not run for a dry_run"),
    ), TestClient(app) as client:
        r = client.post(
            f"/api/convert?{_mint_qs()}",
            files=_form(_banner_docx_bytes(), name="curriculum.docx"),
            data={"dry_run": "1", "markers": _json.dumps(markers)},
        )
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["dry_run"] is True
    assert body["heading1_found"] == 2
    assert [t["title"] for t in body["tabs"]] == ["One", "Two"]
    assert body["retrofit"] == {"markers_matched": 2, "markers_missed": []}
    assert body["problems"] == []
    assert _count_job_rows() == 0


def test_dry_run_with_markers_none_matched_mirrors_failed_envelope():
    """S1-B1 zero-match: the real retrofit fails without converting; the
    plan mirrors that (problems + the retrofit echo with
    candidate_blocks), never a bogus no-headings split plan."""
    import json as _json

    app = _build_app_under_test()

    p1, p2, p3 = _creds_patches()
    with p1, p2, p3, TestClient(app) as client:
        r = client.post(
            f"/api/convert?{_mint_qs()}",
            files=_form(_banner_docx_bytes()),
            data={
                "dry_run": "1",
                "markers": _json.dumps(
                    [{"marker_text": "NO SUCH BANNER", "tab_title": "X"}]
                ),
            },
        )
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["heading1_found"] == 0
    assert body["tabs"] == []
    assert body["split_strategy_used"] == "none"
    assert body["placeholder"] == "none"
    assert body["problems"] and "would FAIL" in body["problems"][0]
    assert body["retrofit"]["markers_matched"] == 0
    missed = body["retrofit"]["markers_missed"]
    assert missed[0]["marker_text"] == "NO SUCH BANNER"
    assert missed[0]["candidate_blocks"], "candidate_blocks aid debugging"
    assert _count_job_rows() == 0


def test_batch_plus_dry_run_is_400():
    app = _build_app_under_test()
    files = [
        ("file", ("a.docx", _docx_bytes(["A"]), _DOCX_MIME)),
        ("file", ("b.docx", _docx_bytes(["B"]), _DOCX_MIME)),
    ]
    with TestClient(app) as client:
        r = client.post(
            "/api/convert", files=files, data={"dry_run": "1"},
            headers=_bearer_headers(),
        )
    assert r.status_code == 400
    assert "dry_run is not supported with batch" in r.json()["error"]


def test_dry_run_drive_file_id_reads_without_write():
    """drive_file_id dry_run: the file is only READ (export/download); the
    converter never runs and no row is created. A .docx drive source gets
    no native-Google-Doc disclosure."""
    from appscriptly.services.drive.api import DOCX_MIME

    app = _build_app_under_test()
    content = _docx_bytes(["Alpha", "Beta"])

    p1, p2, p3 = _creds_patches()
    with p1, p2, p3, patch(
        "appscriptly.preview._fetch_drive_docx_and_mime",
        return_value=(io.BytesIO(content), DOCX_MIME),
    ) as fetch_mock, patch(
        "appscriptly.http_server.routes.convert._convert_docx",
        side_effect=AssertionError("converter must not run for a dry_run"),
    ), TestClient(app) as client:
        r = client.post(
            "/api/convert",
            data={"drive_file_id": "DRIVE123", "dry_run": "1"},
            headers=_bearer_headers(),
        )
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["heading1_found"] == 2
    assert [t["title"] for t in body["tabs"]] == ["Alpha", "Beta"]
    fetch_mock.assert_called_once()
    assert _count_job_rows() == 0
    assert not any("native Google Doc" in note for note in body["info"])


def test_dry_run_drive_gdoc_source_carries_disclosure():
    """S1-M1 residue: a native-Google-Doc drive source behaves differently
    in the real run (copy carries the source's tabs; only the first tab
    splits) - the plan must disclose that."""
    from appscriptly.services.drive.api import GDOC_MIME

    app = _build_app_under_test()
    content = _docx_bytes(["Alpha"])

    p1, p2, p3 = _creds_patches()
    with p1, p2, p3, patch(
        "appscriptly.preview._fetch_drive_docx_and_mime",
        return_value=(io.BytesIO(content), GDOC_MIME),
    ), TestClient(app) as client:
        r = client.post(
            "/api/convert",
            data={"drive_file_id": "GDOC123", "dry_run": "1"},
            headers=_bearer_headers(),
        )
    assert r.status_code == 200, r.text
    notes = r.json()["info"]
    assert any(
        "native Google Doc" in n and "FIRST tab" in n for n in notes
    ), notes


def test_dry_run_malformed_docx_is_400():
    app = _build_app_under_test()
    p1, p2, p3 = _creds_patches()
    with p1, p2, p3, TestClient(app) as client:
        r = client.post(
            f"/api/convert?{_mint_qs()}",
            files=_form(b"PK\x03\x04not-a-real-docx"),
            data={"dry_run": "1"},
        )
    assert r.status_code == 400
    assert "could not read the source as a .docx" in r.json()["error"]


def test_dry_run_invalid_value_is_400():
    app = _build_app_under_test()
    with TestClient(app) as client:
        r = client.post(
            "/api/convert",
            files=_form(_docx_bytes(["A"])),
            data={"dry_run": "maybe"},
            headers=_bearer_headers(),
        )
    assert r.status_code == 400
    assert "dry_run" in r.json()["error"]
