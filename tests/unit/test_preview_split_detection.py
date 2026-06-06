"""Split-strategy detection in preview_tab_split (R2 audit Gap #6).

``preview_tab_split`` is the dry-run users call BEFORE the destructive
.docx -> tabbed-Doc convert. Its ``_detect_split_titles`` helper picks
the tab boundaries per ``split_by`` strategy. The existing preview unit
test (test_preview_threshold_consistency.py) exercises ONLY the
``heading_1`` path + the title-truncation warning. Two strategies were
uncovered:

  * ``page_break`` — walks python-docx runs for the namespace-qualified
    ``<w:br w:type="page"/>`` element via a brittle ``findall``. A typo
    in the namespace or attribute would silently report ZERO page-break
    splits, so a user previewing a page-break-structured doc sees "no
    splits found" and either gives up or runs convert expecting a
    different result. This is exactly the kind of silent, hard-to-notice
    break a dry-run exists to prevent.

  * ``auto`` — cascades heading_1 -> heading_2 -> page_break, returning
    the first strategy that yields any titles (and reporting WHICH one
    it used via ``split_strategy_used``). A regression in the cascade
    order or fall-through would change which structure wins.

Both are pure local logic — no Drive — so they run fast against an
in-memory python-docx Document. ``preview_tab_split`` reads the doc
from ``docx_path`` (no creds needed for the local path).
"""
from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_BREAK

from appscriptly.preview import preview_tab_split


def _save_docx(doc, tmp_path, name: str):
    p = tmp_path / name
    buf = io.BytesIO()
    doc.save(buf)
    p.write_bytes(buf.getvalue())
    return p


def _docx_with_page_breaks(n_breaks: int):
    """A doc with body text split by ``n_breaks`` explicit page breaks.

    python-docx emits ``<w:br w:type="page"/>`` for
    ``run.add_break(WD_BREAK.PAGE)`` — the exact element the detector's
    namespace-qualified findall looks for.
    """
    doc = Document()
    doc.add_paragraph("Intro on page one.")
    for i in range(n_breaks):
        para = doc.add_paragraph()
        run = para.add_run()
        run.add_break(WD_BREAK.PAGE)
        doc.add_paragraph(f"Content after break {i + 1}.")
    return doc


def test_page_break_strategy_detects_each_page_break(tmp_path):
    """``split_by='page_break'`` must find one split per <w:br type=page>.

    Two page breaks => two detected split titles ("Page 2", "Page 3" —
    the detector starts page_idx at 1 and increments before labeling).
    If the namespace-qualified findall regresses, this drops to zero and
    the test fails loudly.
    """
    docx_path = _save_docx(_docx_with_page_breaks(2), tmp_path, "breaks.docx")

    result = preview_tab_split(docx_path=docx_path, split_by="page_break")

    assert result["split_strategy_used"] == "page_break"
    assert result["tab_count"] == 2, (
        f"expected 2 page-break splits, got {result['tab_count']}. "
        f"tabs={result['tabs']!r}. A zero here usually means the "
        f"namespace-qualified <w:br type='page'> findall regressed."
    )
    titles = [t["title"] for t in result["tabs"]]
    assert titles == ["Page 2", "Page 3"]
    # No 'no split points' problem when breaks ARE present.
    assert result["problems"] == []


def test_page_break_strategy_reports_problem_when_no_breaks(tmp_path):
    """A doc with NO page breaks under split_by='page_break' yields zero
    tabs AND the actionable 'No split points found' problem — the
    user-facing signal that the dry-run found nothing to split on."""
    doc = Document()
    doc.add_paragraph("Just one paragraph, no breaks, no headings.")
    docx_path = _save_docx(doc, tmp_path, "flat.docx")

    result = preview_tab_split(docx_path=docx_path, split_by="page_break")

    assert result["tab_count"] == 0
    assert result["problems"], "expected a 'no split points' problem"
    assert "No split points found" in result["problems"][0]
    assert "page_break" in result["problems"][0]


def test_auto_strategy_prefers_heading_1_when_present(tmp_path):
    """``auto`` cascades h1 -> h2 -> page_break and reports which it used.
    With Heading-1s present, it MUST resolve to 'heading_1' (the first
    rung) and return those headings — not fall through to a later rung."""
    doc = Document()
    doc.add_heading("Alpha", level=1)
    doc.add_paragraph("body a")
    doc.add_heading("Beta", level=1)
    doc.add_paragraph("body b")
    docx_path = _save_docx(doc, tmp_path, "h1.docx")

    result = preview_tab_split(docx_path=docx_path, split_by="auto")

    assert result["split_strategy_used"] == "heading_1"
    assert [t["title"] for t in result["tabs"]] == ["Alpha", "Beta"]


def test_auto_strategy_falls_through_to_heading_2_when_no_heading_1(tmp_path):
    """When there are no Heading-1s but there ARE Heading-2s, ``auto``
    must skip the empty h1 rung and resolve to 'heading_2'. Catches a
    cascade that stopped at the first (empty) rung instead of advancing."""
    doc = Document()
    doc.add_heading("Sub One", level=2)
    doc.add_paragraph("body")
    doc.add_heading("Sub Two", level=2)
    doc.add_paragraph("body")
    docx_path = _save_docx(doc, tmp_path, "h2.docx")

    result = preview_tab_split(docx_path=docx_path, split_by="auto")

    assert result["split_strategy_used"] == "heading_2"
    assert [t["title"] for t in result["tabs"]] == ["Sub One", "Sub Two"]


def test_auto_strategy_falls_through_to_page_break_last(tmp_path):
    """With neither Heading-1 nor Heading-2 present, ``auto`` must reach
    the final rung (page_break) and use it when page breaks exist.
    Pins the FULL cascade order h1 -> h2 -> page_break end-to-end."""
    docx_path = _save_docx(_docx_with_page_breaks(1), tmp_path, "auto_pb.docx")

    result = preview_tab_split(docx_path=docx_path, split_by="auto")

    assert result["split_strategy_used"] == "page_break"
    assert result["tab_count"] == 1
    assert result["tabs"][0]["title"] == "Page 2"


def test_auto_strategy_reports_no_splits_when_nothing_matches(tmp_path):
    """``auto`` over a flat doc (no headings, no page breaks) exhausts the
    cascade and reports the 'No split points' problem with strategy
    'auto' — not a stale per-rung strategy name."""
    doc = Document()
    doc.add_paragraph("Flat doc. No structure at all.")
    docx_path = _save_docx(doc, tmp_path, "auto_flat.docx")

    result = preview_tab_split(docx_path=docx_path, split_by="auto")

    assert result["tab_count"] == 0
    assert result["split_strategy_used"] == "auto"
    assert result["problems"]
    assert "No split points found" in result["problems"][0]
