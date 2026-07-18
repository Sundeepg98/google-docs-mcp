"""S1-R1 drift pin: the dry-run plan equals the converter's decision sequence.

``plan_conversion_dry_run`` computes its plan from a LOCAL python-docx
parse adapted into minimal Docs-JSON; the real conversion computes the
same decisions from the Drive-imported document's Docs JSON. Two parse
paths over the same walkers can rot apart silently (the cold review's
S1-B2/M1/M2 were exactly that class), so this suite pins them together:

Every fixture is expressed TWICE - as a real .docx (built with
python-docx) and as the Drive-import-shaped Docs JSON body the real
pipeline would fetch - and both are pushed through to a plan:

  - the .docx through ``plan_conversion_dry_run`` (the adapter path);
  - the Docs JSON through ``_reference_plan``, which replicates the
    plan-relevant decision sequence of ``convert_docx_to_tabbed_doc``
    (steps 2-4: detect / flatten / de-dup against the working copy's
    tabs; step 8: placeholder policy incl. the sole-copy veto; plus the
    no-splits early return) using the converter's OWN seams.

The projections must be EQUAL: strategy, heading1_found, tabs_created,
tab titles + depths, placeholder outcome, veto. Fidelity warnings are
excluded (a documented conservative local subset, disclosed in the
response's info note).

The review's negative control holds here too: perturbing the reference
(e.g. the dedupe seed) makes these fail, so the comparator CAN fail.
"""
from __future__ import annotations

import io
import struct
import zlib

import pytest
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement

from appscriptly.docx_import import (
    _dedupe_split_titles,
    _detect_splits,
    _docapp_children,
    _existing_tab_titles,
    _flatten_splits,
    _unmoved_visible_count,
)
from appscriptly.preview import plan_conversion_dry_run


# ---------------------------------------------------------------------
# Fixture builders - docx side
# ---------------------------------------------------------------------


def _docx_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _png_1px() -> bytes:
    """A valid 1x1 transparent PNG, generated (no fixture file needed)."""
    def chunk(tag: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data)) + tag + data
            + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)
        )
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 6, 0, 0, 0)
    idat = zlib.compress(b"\x00\x00\x00\x00\x00")  # filter 0 + 1 RGBA px
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


# ---------------------------------------------------------------------
# Fixture builders - Drive-shaped Docs JSON side
# ---------------------------------------------------------------------


def _sb() -> dict:
    """The sectionBreak element every real Docs body starts with (the
    REST-only structural record _docapp_children filters out)."""
    return {"sectionBreak": {}}


def _jp(text: str, style: str = "NORMAL_TEXT", *, bullet: bool = False) -> dict:
    para: dict = {
        "paragraphStyle": {"namedStyleType": style},
        "elements": [{"textRun": {"content": f"{text}\n"}}],
    }
    if bullet:
        para["bullet"] = {"listId": "kix.lst0"}
    return {"paragraph": para}


def _jimg() -> dict:
    """An image-only paragraph: a non-textRun inline element and no text -
    visible content per _has_visible_content."""
    return {
        "paragraph": {
            "paragraphStyle": {"namedStyleType": "NORMAL_TEXT"},
            "elements": [
                {"inlineObjectElement": {"inlineObjectId": "kix.img0"}},
                {"textRun": {"content": "\n"}},
            ],
        }
    }


def _jpb() -> dict:
    """A paragraph carrying an explicit page break."""
    return {
        "paragraph": {
            "paragraphStyle": {"namedStyleType": "NORMAL_TEXT"},
            "elements": [{"pageBreak": {}}, {"textRun": {"content": "\n"}}],
        }
    }


def _jtable() -> dict:
    return {"table": {}}


# The working copy's tabs right after a fresh Drive import: exactly one
# tab titled "Tab 1". The reference dedupes against these via the
# converter's own _existing_tab_titles - the same seed the planner must
# mirror (S1-M1).
_FRESH_IMPORT_TABS = [{"tabProperties": {"tabId": "t.0", "title": "Tab 1"}}]


def _reference_plan(
    body_content: list[dict],
    *,
    split_by: str = "heading_1",
    nest_by: str | None = None,
    placeholder_behavior: str = "delete",
) -> dict:
    """The real converter's plan-relevant decisions over Docs JSON.

    Replicates convert_docx_to_tabbed_doc's sequence with the converter's
    own seams: _docapp_children -> _detect_splits -> _flatten_splits ->
    _dedupe_split_titles against _existing_tab_titles(working copy),
    the step-8 placeholder policy with the _unmoved_visible_count veto,
    and the no-splits early return (placeholder "none").
    """
    docapp_children = _docapp_children(body_content)
    splits, strategy = _detect_splits(body_content, split_by, nest_by=nest_by)  # type: ignore[arg-type]
    if not splits:
        return {
            "split_strategy_used": strategy,
            "heading1_found": 0,
            "tabs_created": 0,
            "tabs": [],
            "placeholder": "none",
            "placeholder_veto": None,
        }
    flat = _flatten_splits(splits)
    _dedupe_split_titles(flat, _existing_tab_titles(_FRESH_IMPORT_TABS))

    depths: list[int] = []

    def walk(nodes: list, depth: int) -> None:
        for node in nodes:
            depths.append(depth)
            walk(node["children"], depth + 1)

    walk(splits, 0)

    placeholder = "kept"
    veto: str | None = None
    if placeholder_behavior == "delete":
        all_ranges = [r for s in flat for r in s["ranges"]]
        if _unmoved_visible_count(docapp_children, all_ranges):
            veto = "unmoved_content"
        else:
            placeholder = "deleted"
    elif placeholder_behavior == "rename":
        placeholder = "renamed"

    return {
        "split_strategy_used": strategy,
        "heading1_found": len(splits),
        "tabs_created": len(flat),
        "tabs": [
            {"title": s["title"], "depth": d} for s, d in zip(flat, depths)
        ],
        "placeholder": placeholder,
        "placeholder_veto": veto,
    }


def _projection(plan: dict) -> dict:
    """The equivalence surface of a dry-run plan (fidelity excluded)."""
    return {
        "split_strategy_used": plan["split_strategy_used"],
        "heading1_found": plan["heading1_found"],
        "tabs_created": plan["tabs_created"],
        "tabs": plan["tabs"],
        "placeholder": plan["placeholder"],
        "placeholder_veto": plan.get("placeholder_veto"),
    }


# ---------------------------------------------------------------------
# The parametrized fixtures: (docx builder, docs-json builder, params)
# ---------------------------------------------------------------------


def _case_flat_with_leading_text():
    d = Document()
    d.add_paragraph("Intro before any heading.")
    d.add_heading("Alpha", 1); d.add_paragraph("a body")
    d.add_heading("Beta", 1); d.add_paragraph("b body")
    j = [
        _sb(),
        _jp("Intro before any heading."),
        _jp("Alpha", "HEADING_1"), _jp("a body"),
        _jp("Beta", "HEADING_1"), _jp("b body"),
    ]
    return d, j, {}


def _case_clean_flat_delete():
    d = Document()
    d.add_heading("Alpha", 1); d.add_paragraph("a body")
    d.add_heading("Beta", 1); d.add_paragraph("b body")
    j = [
        _sb(),
        _jp("Alpha", "HEADING_1"), _jp("a body"),
        _jp("Beta", "HEADING_1"), _jp("b body"),
    ]
    return d, j, {}


def _case_image_before_first_h1():
    # The S1-B2 demo: a cover image before the first heading. The veto
    # must SEE it (image-only paragraph = visible content).
    d = Document()
    d.add_picture(io.BytesIO(_png_1px()))
    d.add_heading("Alpha", 1); d.add_paragraph("a body")
    j = [
        _sb(),
        _jimg(),
        _jp("Alpha", "HEADING_1"), _jp("a body"),
    ]
    return d, j, {}


def _case_empty_bullet_before_first_h1():
    # S1-B2's second class: an EMPTY bulleted paragraph is visible via
    # the paragraph-level bullet field.
    d = Document()
    d.add_paragraph("", style="List Bullet")
    d.add_heading("Alpha", 1); d.add_paragraph("a body")
    j = [
        _sb(),
        _jp("", bullet=True),
        _jp("Alpha", "HEADING_1"), _jp("a body"),
    ]
    return d, j, {}


def _case_h1_titled_tab1():
    # The S1-M1 demo: the fresh working copy already holds a tab titled
    # "Tab 1", so an H1 with that exact title must plan as "Tab 1 (2)".
    d = Document()
    d.add_heading("Tab 1", 1); d.add_paragraph("x")
    j = [_sb(), _jp("Tab 1", "HEADING_1"), _jp("x")]
    return d, j, {}


def _case_no_splits():
    # The S1-M2 demo: the real no-splits early return reports
    # placeholder "none".
    d = Document()
    d.add_paragraph("Flat document with no headings at all.")
    j = [_sb(), _jp("Flat document with no headings at all.")]
    return d, j, {}


def _case_nested_orphan_h2():
    # Orphan H2 before the first H1 stays behind (veto); content between
    # an H1 and its first H2 stays in the parent tab.
    d = Document()
    d.add_heading("Orphan", 2)
    d.add_heading("Parent1", 1); d.add_paragraph("between H1 and first H2")
    d.add_heading("Child1a", 2); d.add_paragraph("c1a body")
    d.add_heading("Parent2", 1)
    j = [
        _sb(),
        _jp("Orphan", "HEADING_2"),
        _jp("Parent1", "HEADING_1"), _jp("between H1 and first H2"),
        _jp("Child1a", "HEADING_2"), _jp("c1a body"),
        _jp("Parent2", "HEADING_1"),
    ]
    return d, j, {"nest_by": "heading_2"}


def _case_duplicate_titles():
    d = Document()
    for _ in range(3):
        d.add_heading("Intro", 1); d.add_paragraph("body")
    j = [_sb()] + [
        e for _ in range(3) for e in (_jp("Intro", "HEADING_1"), _jp("body"))
    ]
    return d, j, {}


def _case_empty_heading():
    d = Document()
    d.add_heading("", 1); d.add_paragraph("body")
    j = [_sb(), _jp("", "HEADING_1"), _jp("body")]
    return d, j, {}


def _case_long_title():
    long = "L" * 60
    d = Document()
    d.add_heading(long, 1); d.add_paragraph("body")
    j = [_sb(), _jp(long, "HEADING_1"), _jp("body")]
    return d, j, {}


def _case_table_before_first_h1():
    # A table is an opaque visible block: it vetoes the delete when it
    # precedes the first split, and an H1-styled paragraph INSIDE it
    # must not split (iter_inner_content yields the Table, not its
    # inner paragraphs - matching the JSON side's opaque {"table": {}}).
    d = Document()
    table = d.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].text = "banner text"
    d.add_heading("Alpha", 1); d.add_paragraph("a body")
    j = [
        _sb(),
        _jtable(),
        _jp("Alpha", "HEADING_1"), _jp("a body"),
    ]
    return d, j, {}


def _case_page_break_split():
    d = Document()
    d.add_paragraph("Intro on page one.")
    for i in range(2):
        para = d.add_paragraph()
        para.add_run().add_break(WD_BREAK.PAGE)
        d.add_paragraph(f"Content after break {i + 1}.")
    j = [
        _sb(),
        _jp("Intro on page one."),
        _jpb(), _jp("Content after break 1."),
        _jpb(), _jp("Content after break 2."),
    ]
    return d, j, {"split_by": "page_break"}


def _case_rename_behavior():
    d = Document()
    d.add_heading("Alpha", 1); d.add_paragraph("a body")
    j = [_sb(), _jp("Alpha", "HEADING_1"), _jp("a body")]
    return d, j, {"placeholder_behavior": "rename"}


def _case_equation_before_first_h1():
    # An equation-only paragraph (m:oMath) is visible content: veto.
    d = Document()
    p = d.add_paragraph()
    p._p.append(OxmlElement("m:oMath"))
    d.add_heading("Alpha", 1); d.add_paragraph("a body")
    j = [
        _sb(),
        {
            "paragraph": {
                "paragraphStyle": {"namedStyleType": "NORMAL_TEXT"},
                "elements": [
                    {"equation": {}},
                    {"textRun": {"content": "\n"}},
                ],
            }
        },
        _jp("Alpha", "HEADING_1"), _jp("a body"),
    ]
    return d, j, {}


_CASES = {
    "flat_with_leading_text": _case_flat_with_leading_text,
    "clean_flat_delete": _case_clean_flat_delete,
    "image_before_first_h1": _case_image_before_first_h1,
    "empty_bullet_before_first_h1": _case_empty_bullet_before_first_h1,
    "equation_before_first_h1": _case_equation_before_first_h1,
    "h1_titled_tab1": _case_h1_titled_tab1,
    "no_splits": _case_no_splits,
    "nested_orphan_h2": _case_nested_orphan_h2,
    "duplicate_titles": _case_duplicate_titles,
    "empty_heading": _case_empty_heading,
    "long_title": _case_long_title,
    "table_before_first_h1": _case_table_before_first_h1,
    "page_break_split": _case_page_break_split,
    "rename_behavior": _case_rename_behavior,
}


@pytest.mark.parametrize("case", sorted(_CASES), ids=sorted(_CASES))
def test_plan_equals_converter_decision_sequence(case):
    doc, docs_json, params = _CASES[case]()
    plan = plan_conversion_dry_run(
        docx_bytes=_docx_bytes(doc),
        split_by=params.get("split_by", "heading_1"),
        nest_by=params.get("nest_by"),
        placeholder_behavior=params.get("placeholder_behavior", "delete"),
    )
    reference = _reference_plan(
        docs_json,
        split_by=params.get("split_by", "heading_1"),
        nest_by=params.get("nest_by"),
        placeholder_behavior=params.get("placeholder_behavior", "delete"),
    )
    assert _projection(plan) == reference, (
        f"dry-run plan diverged from the converter's decision sequence "
        f"for case {case!r}"
    )


def test_markers_plan_equals_converter_over_injected_json():
    """The markers path: the plan for a banner doc + markers must equal
    the converter's decisions over the POST-injection Docs JSON (an H1
    inserted before each matched banner - what the real retrofit
    converts)."""
    d = Document()
    d.add_paragraph("MODULE ONE overview")
    d.add_paragraph("module one body")
    d.add_paragraph("MODULE TWO overview")
    d.add_paragraph("module two body")
    markers = [
        {"marker_text": "MODULE ONE", "tab_title": "One"},
        {"marker_text": "MODULE TWO", "tab_title": "Two"},
    ]
    plan = plan_conversion_dry_run(docx_bytes=_docx_bytes(d), markers=markers)

    injected_json = [
        _sb(),
        _jp("One", "HEADING_1"),
        _jp("MODULE ONE overview"), _jp("module one body"),
        _jp("Two", "HEADING_1"),
        _jp("MODULE TWO overview"), _jp("module two body"),
    ]
    assert _projection(plan) == _reference_plan(injected_json)
    assert plan["retrofit"] == {"markers_matched": 2, "markers_missed": []}


# ---------------------------------------------------------------------
# Direct planner tests (the reviewer-noted gap): fidelity scan + veto
# comparator sanity
# ---------------------------------------------------------------------


def test_local_fidelity_scan_reports_equation_via_shared_registry():
    """The conservative local scan renders through the shared
    DROPPED_KINDS registry - an equation in the source produces the
    registry's own warning line."""
    d = Document()
    p = d.add_paragraph()
    p._p.append(OxmlElement("m:oMath"))
    d.add_heading("Alpha", 1); d.add_paragraph("body")
    plan = plan_conversion_dry_run(docx_bytes=_docx_bytes(d))
    assert any("equation" in w and "omitted" in w for w in plan["warnings"])


def test_comparator_can_fail_negative_control():
    """The review's negative control, kept as a permanent fixture: a
    reference computed with a WRONG dedupe seed must NOT equal the plan
    (proves the equivalence assertion can fail)."""
    d = Document()
    d.add_heading("Tab 1", 1); d.add_paragraph("x")
    plan = plan_conversion_dry_run(docx_bytes=_docx_bytes(d))

    docapp = [_sb(), _jp("Tab 1", "HEADING_1"), _jp("x")]
    splits, _ = _detect_splits(docapp, "heading_1")
    flat = _flatten_splits(splits)
    _dedupe_split_titles(flat, set())  # WRONG seed: no "Tab 1"
    wrong_titles = [s["title"] for s in flat]
    assert [t["title"] for t in plan["tabs"]] != wrong_titles
