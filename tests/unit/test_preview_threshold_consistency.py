"""Preview agrees with conversion on the 50-char title limit (unit).

If preview says "fine" but conversion 400s, the dry-run is useless.
This guard makes sure they share the same threshold.

Previously this test lived in tests/integration/ but it doesn't
actually need Drive — preview_tab_split runs locally when given
docx_path. Moved here so it appears in the deploy artifact's
gdocs_test_manifest output as a named regression guard.
"""
from __future__ import annotations

import io


def _docx_with_oversized_heading() -> bytes:
    from docx import Document
    doc = Document()
    long_heading = "A" * 60  # > 50-char API limit
    doc.add_heading(long_heading, level=1)
    doc.add_paragraph("body")
    doc.add_heading("Short", level=1)
    doc.add_paragraph("body2")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_preview_flags_what_convert_truncates(tmp_path):
    """Named regression guard. preview_tab_split MUST warn on titles
    exceeding 50 chars AND truncate them to exactly 50 — matching
    what convert_docx_to_tabbed_doc actually does. If the two
    thresholds drift apart, the dry-run lies.
    """
    from appscriptly.preview import preview_tab_split

    docx_path = tmp_path / "oversize.docx"
    docx_path.write_bytes(_docx_with_oversized_heading())

    result = preview_tab_split(docx_path=docx_path, split_by="heading_1")
    long_tab = next((t for t in result["tabs"] if t["title"] != "Short"), None)
    assert long_tab is not None, "preview didn't detect both headings"
    assert any(
        "truncated" in w.lower() and "50" in w
        for w in long_tab.get("warnings", [])
    ), (
        f"preview did NOT warn about 50-char truncation on long heading. "
        f"warnings={long_tab.get('warnings')!r}. If preview's threshold "
        f"drifts away from convert's actual truncation, dry-run lies."
    )
    # And the truncated title in the preview must indeed be 50 chars.
    assert len(long_tab["title"]) == 50, (
        f"preview title len={len(long_tab['title'])}; should be 50 "
        "(the API limit that convert enforces)."
    )
