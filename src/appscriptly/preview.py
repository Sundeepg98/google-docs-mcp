"""Dry-run preview for the .docx → tabbed Google Doc conversion.

Reads a .docx locally (no Drive operations) and reports what the tab
split would look like, plus any validation problems. Useful for catching
missing headings or over-length titles BEFORE creating a Google Doc.

For inputs already on Drive (raw .docx OR native Google Doc), we
download/export as .docx bytes via Drive, then parse the same way —
keeps the validation logic uniform.

``plan_conversion_dry_run`` (the ``/api/convert?dry_run=1`` engine) goes
further than ``preview_tab_split``: it adapts the local python-docx parse
into the minimal Google-Docs element shape the REAL converter walks, then
runs the converter's OWN split/nest/placeholder logic
(``docx_import._detect_splits`` et al.) over it, so the predicted plan is
computed by the same code the conversion uses — no forked re-implement.
A ``markers`` request likewise reuses retrofit's own local injection
(``retrofit._inject_headings``) before planning. Still ZERO Drive writes:
the upload path never touches Drive, and the drive_file_id path only
reads/exports bytes. The drift pin
(``tests/unit/test_dry_run_plan_equivalence.py``) holds this path equal to
the converter's decision sequence over Drive-shaped Docs JSON.
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import Literal, cast
from zipfile import BadZipFile

from docx import Document
from docx.document import Document as DocumentT  # the class; `Document` itself is a factory function
from docx.opc.exceptions import PackageNotFoundError
from google.oauth2.credentials import Credentials
from appscriptly.google_api_client import execute_with_retry
from appscriptly.google_clients import get_service

from .services.drive.api import DOCX_MIME, GDOC_MIME

# The dry-run planner reuses the converter's OWN split/nest/placeholder
# machinery (so the predicted plan can never drift from what the real
# conversion does) plus the shared fidelity registry. These are the pure
# planning seams the wave-3 contract names for reuse; none of them make a
# Drive call. preview -> docx_import is a one-way import (docx_import does
# not import preview), so there is no cycle.
from .docx_import import (
    _MAX_TAB_TITLE,
    NestBy,
    SplitBy,
    _dedupe_split_titles,
    _detect_splits,
    _docapp_children,
    _extract_paragraph_text,
    _flatten_splits,
    _SplitPoint,
    _unmoved_visible_count,
)
from .retrofit import _inject_headings
from .services.docs.content_transplant import FidelityReport

PreviewSplitBy = Literal["heading_1", "heading_2", "page_break", "auto"]

TITLE_MAX_CHARS = 50  # Google Docs API hard limit (returns 400 above this)

_STYLE_FOR_SPLIT = {
    "heading_1": "Heading 1",
    "heading_2": "Heading 2",
}


def preview_tab_split(
    creds: Credentials | None = None,
    docx_path: Path | None = None,
    drive_file_id: str | None = None,
    split_by: PreviewSplitBy = "heading_1",
) -> dict:
    """Return what tabs would be created without actually creating anything.

    Provide exactly one of ``docx_path`` or ``drive_file_id``. For
    ``drive_file_id``, ``creds`` is required (we need to fetch the
    bytes); for ``docx_path`` it can be omitted.

    Returns ``{"split_strategy_used", "tab_count", "tabs":
    [{"title", "warnings": [...]}, ...], "problems": [...]}``.
    """
    if (docx_path is None) == (drive_file_id is None):
        raise ValueError(
            "Provide exactly one of docx_path or drive_file_id."
        )

    if docx_path is not None:
        if not docx_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")
        doc = Document(str(docx_path))
    else:
        if creds is None:
            raise ValueError("creds required when previewing a Drive file")
        buf = _fetch_drive_as_docx(creds, drive_file_id)  # type: ignore[arg-type]
        doc = Document(buf)

    detected_titles, strategy_used = _detect_split_titles(doc, split_by)
    problems: list[str] = []
    if not detected_titles:
        problems.append(
            f"No split points found with strategy '{split_by}'. The .docx "
            "has no paragraphs matching the expected heading style. "
            "Either change split_by or inject markers via retrofit_existing_docx."
        )

    tabs: list[dict] = []
    for raw_title in detected_titles:
        truncated = raw_title[:TITLE_MAX_CHARS].strip()
        warnings: list[str] = []
        if len(raw_title) > TITLE_MAX_CHARS:
            warnings.append(
                f"title is {len(raw_title)} chars; will be truncated to {TITLE_MAX_CHARS}"
            )
        if not truncated:
            warnings.append("title is empty after stripping; will fall back to 'Section N'")
        tabs.append({"title": truncated or None, "raw_title": raw_title, "warnings": warnings})

    return {
        "split_strategy_used": strategy_used,
        "tab_count": len(tabs),
        "tabs": tabs,
        "problems": problems,
    }


def _detect_split_titles(
    doc: DocumentT, split_by: PreviewSplitBy
) -> tuple[list[str], str]:
    """Walk the python-docx Document and emit titles per split strategy."""
    if split_by == "auto":
        for strategy in ("heading_1", "heading_2", "page_break"):
            titles, _ = _detect_split_titles(doc, strategy)  # type: ignore[arg-type]
            if titles:
                return titles, strategy
        return [], "auto"

    target_style = _STYLE_FOR_SPLIT.get(split_by)
    titles: list[str] = []

    if split_by in ("heading_1", "heading_2"):
        for para in doc.paragraphs:
            if para.style and para.style.name == target_style:
                titles.append(para.text.strip())
    elif split_by == "page_break":
        # python-docx exposes page breaks via runs containing <w:br type="page"/>
        page_idx = 1
        for para in doc.paragraphs:
            for run in para.runs:
                if run._element.findall(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br[@{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type='page']"
                ):
                    page_idx += 1
                    titles.append(f"Page {page_idx}")

    return titles, split_by


def _fetch_drive_as_docx(creds: Credentials, drive_file_id: str) -> io.BytesIO:
    """Download a Drive file as .docx bytes (works for both .docx and Google Doc).

    PR-Δ3.5: helper for ``preview_tab_split`` which is annotated
    ``readonly=True``. All three Drive calls (metadata fetch, raw
    media download, .docx export) wrapped with idempotent retry.
    """
    buf, _ = _fetch_drive_docx_and_mime(creds, drive_file_id)
    return buf


def _fetch_drive_docx_and_mime(
    creds: Credentials, drive_file_id: str
) -> tuple[io.BytesIO, str]:
    """As ``_fetch_drive_as_docx``, but also return the source mimeType.

    The dry-run planner needs it: a native Google Doc source (GDOC_MIME)
    behaves differently in the real conversion (copy carries the source's
    tabs; only the first tab is split), which the plan must disclose.
    """
    drive = get_service("drive", "v3", credentials=creds)
    meta = execute_with_retry(
        lambda: drive.files().get(fileId=drive_file_id, fields="mimeType").execute(),
        idempotent=True,
        op_name="drive.files.get.preview_meta",
    )
    mime = meta.get("mimeType")
    if mime == DOCX_MIME:
        buf = execute_with_retry(
            lambda: drive.files().get_media(fileId=drive_file_id).execute(),
            idempotent=True,
            op_name="drive.files.get_media.docx",
        )
        return io.BytesIO(buf), mime
    if mime == GDOC_MIME:
        buf = execute_with_retry(
            lambda: drive.files().export(
                fileId=drive_file_id, mimeType=DOCX_MIME
            ).execute(),
            idempotent=True,
            op_name="drive.files.export.docx",
        )
        return io.BytesIO(buf), mime
    raise ValueError(
        f"Drive file {drive_file_id!r} has mimeType {mime!r}. "
        "Expected .docx or Google Doc."
    )


# ---------------------------------------------------------------------
# Dry-run conversion plan (the /api/convert?dry_run=1 engine)
# ---------------------------------------------------------------------
#
# The upload path holds only .docx bytes with no Google Doc, so it CANNOT
# obtain the Docs-JSON body ``_detect_splits`` / ``scan_source_fidelity``
# walk without importing to Drive first (a write). To stay ZERO-write we
# parse the .docx locally and adapt it into the MINIMAL element shape
# those seams read - just enough for the split/nest/placeholder walk -
# then hand it to the converter's own functions. The predicted plan is
# therefore produced by the SAME code the real conversion runs; a local
# preview cannot promise Drive's exact docx->Doc rendering, but the split
# structure it derives is the converter's, not a fork.

# Word / OOXML namespaces used by the conservative local fidelity scan.
_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_M_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"

# python-docx style name -> Docs namedStyleType. Only HEADING_1 / HEADING_2
# actually drive splits; everything else maps to NORMAL_TEXT (its exact
# value is irrelevant to _detect_splits, which only compares against the
# heading style for the chosen strategy).
_STYLE_NAME_TO_NAMED_TYPE = {
    "Heading 1": "HEADING_1",
    "Heading 2": "HEADING_2",
    "Heading 3": "HEADING_3",
    "Heading 4": "HEADING_4",
    "Heading 5": "HEADING_5",
    "Heading 6": "HEADING_6",
    "Title": "TITLE",
    "Subtitle": "SUBTITLE",
}


def _para_has_page_break(para) -> bool:
    """True when a python-docx paragraph carries an explicit page break
    (the ``<w:br w:type="page"/>`` element the page_break split walks)."""
    for run in para.runs:
        if run._element.findall(
            f".//{_W_NS}br[@{_W_NS}type='page']"
        ):
            return True
    return False


# Non-text inline content the placeholder-veto scan must SEE (S1-B2):
# images (w:drawing), legacy pictures (w:pict), equations (m:oMath). In
# real Docs JSON such paragraphs carry a non-textRun element, which
# ``_has_visible_content`` counts as visible even when the paragraph has
# no text - so an image-only paragraph before the first heading vetoes
# the placeholder delete. The adapter must not render them invisible.
_NONTEXT_PARA_TAGS = frozenset({
    f"{_W_NS}drawing",
    f"{_W_NS}pict",
    f"{_M_NS}oMath",
})


def _para_has_nontext_content(para) -> bool:
    """True when the paragraph XML carries an image / picture / equation."""
    for node in para._p.iter():
        if node.tag in _NONTEXT_PARA_TAGS:
            return True
    return False


def _para_is_list_item(para) -> bool:
    """True for list paragraphs (direct numbering or a List* style).

    The real Docs JSON gives list paragraphs a ``bullet`` field, which
    counts as visible content even when the text is empty - an empty
    bulleted paragraph before the first heading vetoes the placeholder
    delete, so the adapter must carry the signal (S1-B2)."""
    pPr = para._p.pPr
    if pPr is not None and pPr.numPr is not None:
        return True
    style = para.style
    return bool(style is not None and (style.name or "").startswith("List"))


def _docx_to_min_body(doc: DocumentT) -> list[dict]:
    """Adapt a python-docx Document into the minimal Docs-``body.content``
    shape ``_detect_splits`` / ``_docapp_children`` / ``_unmoved_visible_count``
    read.

    Only the fields those walkers touch are emitted:
      - each body paragraph -> ``{"paragraph": {"paragraphStyle":
        {"namedStyleType": ...}, "elements": [{"textRun": {"content":
        text}}]}}``, plus ``{"pageBreak": {}}`` when the paragraph breaks
        a page, ``{"inlineObjectElement": {}}`` when it carries an image /
        picture / equation (S1-B2 visibility), and a paragraph-level
        ``"bullet"`` key for list paragraphs;
      - each body table -> ``{"table": {}}`` - an opaque non-paragraph
        block, so it extends the current split range and counts as
        VISIBLE content for the placeholder-veto check, exactly as a real
        table does in the converted Doc.

    Document order is preserved via ``iter_inner_content`` so the emitted
    index space matches the walk order the converter uses.
    """
    body: list[dict] = []
    for block in doc.iter_inner_content():
        # Table (or any non-paragraph block): opaque, visible, non-splitting.
        if not hasattr(block, "runs"):
            body.append({"table": {}})
            continue
        style_name = block.style.name if block.style is not None else None
        named_type = _STYLE_NAME_TO_NAMED_TYPE.get(style_name or "", "NORMAL_TEXT")
        elements: list[dict] = [{"textRun": {"content": block.text}}]
        if _para_has_nontext_content(block):
            elements.append({"inlineObjectElement": {}})
        if _para_has_page_break(block):
            elements.append({"pageBreak": {}})
        para: dict = {
            "paragraphStyle": {"namedStyleType": named_type},
            "elements": elements,
        }
        if _para_is_list_item(block):
            # Non-empty like the real shape ({"listId": ...}): the veto
            # scan tests the field with bool(), so {} would vanish.
            para["bullet"] = {"listId": "local"}
        body.append({"paragraph": para})
    return body


def _scan_local_fidelity(doc: DocumentT) -> FidelityReport:
    """Conservative content-fidelity scan of the SOURCE .docx, reusing the
    shared ``FidelityReport`` + registry.

    Only kinds that are (a) unambiguous in the .docx XML and (b) always
    dropped/degraded by the transplant are counted - equations and
    footnotes have no REST write path (DROPPED), a table of contents is
    replaced with a placeholder (DEGRADED). Ambiguous kinds (e.g. inline
    vs. floating drawings) are deliberately NOT scanned locally: a false
    fidelity warning is worse than none. The AUTHORITATIVE report is the
    one the real conversion returns after Drive renders the Doc.
    """
    report = FidelityReport()
    equations = footnotes = toc = 0
    for el in doc.element.body.iter():
        tag = el.tag
        if tag == f"{_M_NS}oMath":
            equations += 1
        elif tag == f"{_W_NS}footnoteReference":
            footnotes += 1
        elif tag == f"{_W_NS}fldSimple":
            if "TOC" in (el.get(f"{_W_NS}instr") or "").upper():
                toc = 1
        elif tag == f"{_W_NS}instrText":
            if "TOC" in (el.text or "").upper():
                toc = 1
    report.count("equation", equations)
    report.count("footnote", footnotes)
    report.count("toc", toc)
    return report


def _load_dry_run_doc(
    creds: Credentials | None,
    docx_bytes: bytes | None,
    docx_path: Path | None,
    drive_file_id: str | None,
) -> tuple[DocumentT, str | None]:
    """Resolve exactly one input mode to ``(Document, source_mime)`` with
    NO Drive write (the drive path only reads/exports bytes).

    ``source_mime`` is the Drive mimeType for a drive_file_id input and
    None otherwise - the planner discloses native-Google-Doc specifics."""
    provided = [x is not None for x in (docx_bytes, docx_path, drive_file_id)]
    if sum(provided) != 1:
        raise ValueError(
            "Provide exactly one of docx_bytes, docx_path, or drive_file_id."
        )
    # A truncated / non-OOXML payload surfaces as BadZipFile (bad bytes) or
    # PackageNotFoundError (valid zip, not an OPC package). Normalize both to
    # a ValueError so the caller answers 400, not a bare 500.
    try:
        if docx_bytes is not None:
            return Document(io.BytesIO(docx_bytes)), None
        if docx_path is not None:
            if not docx_path.exists():
                raise FileNotFoundError(f"DOCX file not found: {docx_path}")
            return Document(str(docx_path)), None
        if creds is None:
            raise ValueError("creds required when previewing a Drive file")
        buf, mime = _fetch_drive_docx_and_mime(creds, drive_file_id)  # type: ignore[arg-type]
        return Document(buf), mime
    except (BadZipFile, PackageNotFoundError) as e:
        raise ValueError(
            f"could not read the source as a .docx ({e}); ensure it is a "
            "valid .docx document"
        ) from e


def _split_depths(splits: list[_SplitPoint]) -> list[int]:
    """Pre-order depth per split node, aligned 1:1 with ``_flatten_splits``
    (0 = top-level parent, 1 = ``nest_by`` child)."""
    out: list[int] = []

    def walk(nodes: list[_SplitPoint], depth: int) -> None:
        for node in nodes:
            out.append(depth)
            walk(node["children"], depth + 1)

    walk(splits, 0)
    return out


# The working copy of a FRESH Drive import carries exactly one tab,
# titled "Tab 1" - the set the real converter's title de-dup runs against
# (docx_import.py step 4: _dedupe_split_titles over _existing_tab_titles
# of the just-imported doc). The plan must seed the same set or an H1
# literally titled "Tab 1" plans as "Tab 1" while the real run produces
# "Tab 1 (2)" (S1-M1).
_FRESH_IMPORT_TAB_TITLES = frozenset({"Tab 1"})

# Disclosure for a native-Google-Doc drive source: the local parse sees
# the EXPORTED docx, while the real conversion copies the doc - carrying
# ALL its tabs (extra de-dup surface) and splitting only the first tab.
_GDOC_SOURCE_INFO = (
    "drive_file_id source is a native Google Doc: this plan derives from "
    "its exported .docx content. The real conversion copies the document, "
    "splits only the FIRST tab, and carries the source's other tabs into "
    "the working copy, which can add title de-duplication suffixes not "
    "shown here."
)

# Every dry-run response carries this honesty note: nothing was created,
# the URL is intact, and BOTH the split plan and the fidelity warnings
# come from a local parse of the source - the conversion itself is the
# authority.
_DRY_RUN_INFO_NOTE = (
    "dry_run: no document was created and no signed URL was consumed. "
    "POST the identical request without dry_run to perform the "
    "conversion. The split plan and the content-fidelity warnings come "
    "from a local parse of the source document; the authoritative result "
    "comes from the conversion itself."
)


def plan_conversion_dry_run(
    creds: Credentials | None = None,
    *,
    docx_bytes: bytes | None = None,
    docx_path: Path | None = None,
    drive_file_id: str | None = None,
    split_by: str = "heading_1",
    nest_by: str | None = None,
    placeholder_behavior: str = "delete",
    markers: list[dict] | None = None,
) -> dict:
    """Return the conversion PLAN for a .docx without creating anything.

    Provide exactly one input: ``docx_bytes`` (the /api/convert upload
    path), ``docx_path`` (local), or ``drive_file_id`` (``creds`` required;
    read/export only). ``nest_by``, ``placeholder_behavior`` and
    ``markers`` are the same values ``convert_endpoint`` validated.

    ``markers`` mirrors the retrofit path: the SAME local injection the
    real ``retrofit_existing_docx`` performs (``_inject_headings``) runs
    on the parsed document first, the plan reflects the injected Heading
    1s, and the response carries the real path's ``retrofit`` echo
    (``markers_matched`` / ``markers_missed``). When nothing matches, the
    real run returns a failed envelope without converting - the plan
    mirrors that as a ``problems`` entry instead of a bogus split plan.

    The response mirrors the REAL convert result's echo field NAMES so a
    caller can line up plan vs. outcome, plus ``dry_run: true``:
    ``{dry_run, split_strategy_used, heading1_found, tabs_created, tabs:
    [{title, depth}, ...] (pre-order), placeholder, placeholder_veto?,
    warnings, info, problems, retrofit?}``. ``heading1_found`` counts
    top-level splits (parents); ``tabs_created`` includes ``nest_by``
    children - exactly as the converter reports them.
    """
    doc, source_mime = _load_dry_run_doc(creds, docx_bytes, docx_path, drive_file_id)

    warnings: list[str] = []
    info: list[str] = []
    problems: list[str] = []

    # markers (S1-B1): run retrofit's OWN local injection before planning,
    # so the plan reflects the same synthetic Heading 1s the real run
    # converts on. Zero-match mirrors the real failed envelope.
    retrofit_echo: dict | None = None
    if markers:
        matched, missed_specs = _inject_headings(doc, markers)
        retrofit_echo = {
            "markers_matched": matched,
            "markers_missed": missed_specs,
        }
        if not matched:
            problems.append(
                "None of the marker_text values matched any block in the "
                "document; the real conversion would FAIL for this request "
                "without creating anything. Check the candidate_blocks "
                "list under retrofit.markers_missed for the actual visible "
                "text of each body block, pick a distinctive substring, "
                "and retry."
            )
            info.append(_DRY_RUN_INFO_NOTE)
            return {
                "dry_run": True,
                "split_strategy_used": "none",
                "heading1_found": 0,
                "tabs_created": 0,
                "tabs": [],
                "placeholder": "none",
                "warnings": warnings,
                "info": info,
                "problems": problems,
                "retrofit": retrofit_echo,
            }

    # Split / nest / placeholder: computed by the converter's own walkers
    # over the adapted body, so the plan cannot drift from the real run.
    # nest_by reaches here already validated by convert_endpoint (None or
    # "heading_2"); narrow to the NestBy literal for the walker.
    nest: NestBy | None = "heading_2" if nest_by == "heading_2" else None
    body_content = _docx_to_min_body(doc)
    docapp_children = _docapp_children(body_content)
    # split_by is validated by convert_endpoint against the SplitBy set.
    splits, strategy_used = _detect_splits(
        body_content, cast(SplitBy, split_by), nest_by=nest
    )
    flat_splits = _flatten_splits(splits)
    # The converter's same-title de-dup, seeded exactly like the real run:
    # a fresh working copy's only tab is "Tab 1" (S1-M1), and duplicate
    # H1s become "Intro" / "Intro (2)".
    _dedupe_split_titles(flat_splits, set(_FRESH_IMPORT_TAB_TITLES))

    # Per-split title diagnostics. A split's first range starts at its
    # heading paragraph's docapp index, so we can recover the RAW
    # (pre-truncation) heading text to warn about truncation - the
    # converter truncates silently; the dry-run makes it visible.
    tabs: list[dict] = []
    depths = _split_depths(splits)
    for split, depth in zip(flat_splits, depths):
        tabs.append({"title": split["title"], "depth": depth})
        if strategy_used in ("heading_1", "heading_2"):
            raw_idx = split["ranges"][0][0]
            raw_title = _extract_paragraph_text(
                docapp_children[raw_idx].get("paragraph", {})
            )
            if len(raw_title) > _MAX_TAB_TITLE:
                warnings.append(
                    f"tab title {raw_title[:20]!r}... is {len(raw_title)} "
                    f"characters; it will be truncated to {_MAX_TAB_TITLE} "
                    f"(final title {split['title']!r})"
                )
            elif not raw_title:
                warnings.append(
                    f"a {strategy_used} heading has no text; its tab falls "
                    f"back to {split['title']!r}"
                )

    if not splits:
        problems.append(
            f"No split points found with strategy {split_by!r}. The document "
            "has no paragraphs matching the expected heading style; the "
            "conversion would leave it as a single-tab import. Change "
            "split_by, or inject Heading 1s via the markers/retrofit path."
        )

    # Fidelity: reuse the shared registry (DROPPED -> warnings, DEGRADED ->
    # notes). Conservative local subset; see _scan_local_fidelity.
    report = _scan_local_fidelity(doc)
    warnings.extend(report.warnings)
    info.extend(report.notes)

    # Placeholder decision, mirroring convert_docx_to_tabbed_doc step 8:
    # a "delete" is VETOED when visible content sits before the first
    # split point (it would be the only copy). rename/keep are honored.
    # NO splits mirrors the real no-splits early return, which reports
    # "none" (the single-tab import has no placeholder duplication) -
    # S1-M2.
    placeholder_outcome = "kept" if splits else "none"
    placeholder_veto: str | None = None
    if splits and placeholder_behavior == "delete":
        all_ranges = [r for s in flat_splits for r in s["ranges"]]
        if _unmoved_visible_count(docapp_children, all_ranges):
            placeholder_veto = "unmoved_content"
            warnings.append(
                "placeholder tab would be KEPT instead of deleted: content "
                "before the first split point is never moved into a tab, and "
                "deleting the tab would destroy the only copy."
            )
        else:
            placeholder_outcome = "deleted"
    elif splits and placeholder_behavior == "rename":
        placeholder_outcome = "renamed"

    if source_mime == GDOC_MIME:
        info.append(_GDOC_SOURCE_INFO)
    info.append(_DRY_RUN_INFO_NOTE)

    result: dict = {
        "dry_run": True,
        "split_strategy_used": strategy_used,
        "heading1_found": len(splits),
        "tabs_created": len(flat_splits),
        "tabs": tabs,
        "placeholder": placeholder_outcome,
        "warnings": warnings,
        "info": info,
        "problems": problems,
    }
    if placeholder_veto:
        result["placeholder_veto"] = placeholder_veto
    if retrofit_echo is not None:
        result["retrofit"] = retrofit_echo
    return result
