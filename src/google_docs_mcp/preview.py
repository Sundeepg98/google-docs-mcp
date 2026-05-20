"""Dry-run preview for the .docx → tabbed Google Doc conversion.

Reads a .docx locally (no Drive operations) and reports what the tab
split would look like, plus any validation problems. Useful for catching
missing headings or over-length titles BEFORE creating a Google Doc.

For inputs already on Drive (raw .docx OR native Google Doc), we
download/export as .docx bytes via Drive, then parse the same way —
keeps the validation logic uniform.
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import Literal

from docx import Document
from docx.document import Document as DocumentT  # the class; `Document` itself is a factory function
from google.oauth2.credentials import Credentials
from google_docs_mcp.google_clients import get_service

from .services.drive.api import DOCX_MIME, GDOC_MIME

PreviewSplitBy = Literal["heading_1", "heading_2", "page_break", "auto"]

TITLE_MAX_CHARS = 50  # Google Docs API hard limit (returns 400 above this)

_STYLE_FOR_SPLIT = {
    "heading_1": "Heading 1",
    "heading_2": "Heading 2",
}


def preview_tab_split(
    creds: Credentials | None = None,
    docx_path: Path | None = None,
    drive_file_id: str | None = None,
    split_by: PreviewSplitBy = "heading_1",
) -> dict:
    """Return what tabs would be created without actually creating anything.

    Provide exactly one of ``docx_path`` or ``drive_file_id``. For
    ``drive_file_id``, ``creds`` is required (we need to fetch the
    bytes); for ``docx_path`` it can be omitted.

    Returns ``{"split_strategy_used", "tab_count", "tabs":
    [{"title", "warnings": [...]}, ...], "problems": [...]}``.
    """
    if (docx_path is None) == (drive_file_id is None):
        raise ValueError(
            "Provide exactly one of docx_path or drive_file_id."
        )

    if docx_path is not None:
        if not docx_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")
        doc = Document(str(docx_path))
    else:
        if creds is None:
            raise ValueError("creds required when previewing a Drive file")
        buf = _fetch_drive_as_docx(creds, drive_file_id)  # type: ignore[arg-type]
        doc = Document(buf)

    detected_titles, strategy_used = _detect_split_titles(doc, split_by)
    problems: list[str] = []
    if not detected_titles:
        problems.append(
            f"No split points found with strategy '{split_by}'. The .docx "
            "has no paragraphs matching the expected heading style. "
            "Either change split_by or inject markers via retrofit_existing_docx."
        )

    tabs: list[dict] = []
    for raw_title in detected_titles:
        truncated = raw_title[:TITLE_MAX_CHARS].strip()
        warnings: list[str] = []
        if len(raw_title) > TITLE_MAX_CHARS:
            warnings.append(
                f"title is {len(raw_title)} chars; will be truncated to {TITLE_MAX_CHARS}"
            )
        if not truncated:
            warnings.append("title is empty after stripping; will fall back to 'Section N'")
        tabs.append({"title": truncated or None, "raw_title": raw_title, "warnings": warnings})

    return {
        "split_strategy_used": strategy_used,
        "tab_count": len(tabs),
        "tabs": tabs,
        "problems": problems,
    }


def _detect_split_titles(
    doc: DocumentT, split_by: PreviewSplitBy
) -> tuple[list[str], str]:
    """Walk the python-docx Document and emit titles per split strategy."""
    if split_by == "auto":
        for strategy in ("heading_1", "heading_2", "page_break"):
            titles, _ = _detect_split_titles(doc, strategy)  # type: ignore[arg-type]
            if titles:
                return titles, strategy
        return [], "auto"

    target_style = _STYLE_FOR_SPLIT.get(split_by)
    titles: list[str] = []

    if split_by in ("heading_1", "heading_2"):
        for para in doc.paragraphs:
            if para.style and para.style.name == target_style:
                titles.append(para.text.strip())
    elif split_by == "page_break":
        # python-docx exposes page breaks via runs containing <w:br type="page"/>
        page_idx = 1
        for para in doc.paragraphs:
            for run in para.runs:
                if run._element.findall(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br[@{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type='page']"
                ):
                    page_idx += 1
                    titles.append(f"Page {page_idx}")

    return titles, split_by


def _fetch_drive_as_docx(creds: Credentials, drive_file_id: str) -> io.BytesIO:
    """Download a Drive file as .docx bytes (works for both .docx and Google Doc)."""
    drive = get_service("drive", "v3", credentials=creds)
    meta = drive.files().get(fileId=drive_file_id, fields="mimeType").execute()
    mime = meta.get("mimeType")
    if mime == DOCX_MIME:
        buf = drive.files().get_media(fileId=drive_file_id).execute()
        return io.BytesIO(buf)
    if mime == GDOC_MIME:
        buf = drive.files().export(
            fileId=drive_file_id, mimeType=DOCX_MIME
        ).execute()
        return io.BytesIO(buf)
    raise ValueError(
        f"Drive file {drive_file_id!r} has mimeType {mime!r}. "
        "Expected .docx or Google Doc."
    )
